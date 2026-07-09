import streamlit as st
import docx
from docx.shared import RGBColor
import language_tool_python
import re
from io import BytesIO

# 初始化语法检查引擎
@st.cache_resource
def load_tool():
    return language_tool_python.LanguageTool('en-US')

tool = load_tool()

# 精准提取并切分出纯英文片段
def get_english_chunks(text):
    # 正则表达式：匹配连续的英文单词、数字、英文标点和空格
    # 这会把中文、中文标点完全排除在外
    return list(re.finditer(r'[a-zA-Z0-9\s\.,!\?\'"\(\)\-\;:—’“”‘’]+', text))

# --- 网页界面设计 ---
st.set_page_config(page_title="双语剧本纯英文精准校对工具", page_icon="🎬", layout="centered")

st.title("🎬 双语剧本纯英文精准校对工具")
st.write("上传你的中英双语剧本。系统会自动无视所有中文，【只识别英文】是否有误。错误地方将标红并附带修改建议。")

st.markdown("---")

uploaded_file = st.file_uploader("请上传你的双语剧本 (.docx)", type=["docx"])

if uploaded_file is not None:
    st.info("正在精准扫描英文语法，请稍候...")
    
    try:
        doc = docx.Document(uploaded_file)
        new_doc = docx.Document()
        
        total_errors = 0
        
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            
            # 如果是空行，直接保留
            if not original_text.strip():
                new_doc.add_paragraph("")
                continue
                
            # 建立一个新段落，并保持原有的对齐格式
            new_para = new_doc.add_paragraph()
            new_para.alignment = paragraph.alignment
            
            # 找出当前行所有的英文片段
            chunks = get_english_chunks(original_text)
            
            if not chunks:
                # 如果这一行完全没有英文（全是中文），直接原样写入
                new_para.add_run(original_text)
                continue
                
            last_text_idx = 0
            
            for chunk in chunks:
                chunk_start = chunk.start()
                chunk_end = chunk.end()
                chunk_text = chunk.group()
                
                # 1. 写入英文片段前的中文或多余字符
                if chunk_start > last_text_idx:
                    new_para.add_run(original_text[last_text_idx:chunk_start])
                
                # 如果这个英文片段全是空格，直接写入，不检查
                if not chunk_text.strip():
                    new_para.add_run(chunk_text)
                    last_text_idx = chunk_end
                    continue
                
                # 2. 对这一小段英文进行语法检查
                matches = tool.check(chunk_text)
                
                if not matches:
                    # 这段英文完全正确，原样写入
                    new_para.add_run(chunk_text)
                else:
                    total_errors += len(matches)
                    last_chunk_idx = 0
                    
                    # 在这个英文片段内部进行更细致的错词定位标红
                    for match in matches:
                        start = match.offset
                        end = match.offset + match.error_length
                        
                        # 写入英文片段内部错词前的正常英文
                        if start > last_chunk_idx:
                            new_para.add_run(chunk_text[last_chunk_idx:start])
                            
                        # 写入错误的英文，并将其【标红】
                        error_run = new_para.add_run(chunk_text[start:end])
                        error_run.font.color.rgb = RGBColor(255, 0, 0)
                        error_run.bold = True
                        
                        # 紧跟其后插入【绿色的修改建议】
                        if match.replacements:
                            suggestion = match.replacements[0]
                            suggest_run = new_para.add_run(f"【👉 建议改为: {suggestion}】")
                            suggest_run.font.color.rgb = RGBColor(0, 128, 0)
                            suggest_run.bold = True
                            
                        last_chunk_idx = end
                        
                    # 写入英文片段内部尾部的剩余英文
                    if last_chunk_idx < len(chunk_text):
                        new_para.add_run(chunk_text[last_chunk_idx:])
                        
                last_text_idx = chunk_end
                
            # 3. 写入整行末尾可能遗留的中文或标点
            if last_text_idx < len(original_text):
                new_para.add_run(original_text[last_text_idx:])
                    
        # 保存到内存提供下载
        output = BytesIO()
        new_doc.save(output)
        output.seek(0)
        
        st.balloons()
        if total_errors > 0:
            st.success(f"✨ 校对完成！全篇共发现 {total_errors} 处英文语法/拼写可疑点，中文已自动略过。")
        else:
            st.success("🎉 全篇非常完美！英文部分未检测到任何明显的语法错误。")
            
        st.download_button(
            label="📥 下载【纯英文校对标红版】剧本 (.docx)",
            data=output,
            file_name=f"Verified_{uploaded_file.name}",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except Exception as e:
        st.error(f"处理文档时出错: {e}")
